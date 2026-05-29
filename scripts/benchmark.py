"""Benchmark DocxMint vs python-docx across document sizes.

Measures wall-clock time to build and save a document with N paragraphs
and an N/10-row table. Results are written to scripts/benchmark_results.json
for use by scripts/generate_charts.py.

Usage:
    python scripts/benchmark.py
"""

from __future__ import annotations

import json
import os
import statistics
import tempfile
import time

DOCUMENT_SIZES = [100, 500, 1_000, 5_000, 10_000]

def _repeats(n: int) -> int:
    if n >= 5_000:
        return 2
    if n >= 1_000:
        return 3
    return 5


def _temp_path(suffix: str) -> str:
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    return path


# ---------------------------------------------------------------------------
# python-docx
# ---------------------------------------------------------------------------

def bench_python_docx(n_paragraphs: int) -> float:
    import docx

    times = []
    for _ in range(_repeats(n_paragraphs)):
        path = _temp_path(".docx")
        try:
            t0 = time.perf_counter()
            doc = docx.Document()
            for i in range(n_paragraphs):
                doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
            n_rows = max(1, n_paragraphs // 10)
            table = doc.add_table(rows=n_rows, cols=3)
            for r in range(n_rows):
                table.cell(r, 0).text = f"Row {r}"
                table.cell(r, 1).text = "Value A"
                table.cell(r, 2).text = "Value B"
            doc.save(path)
            times.append(time.perf_counter() - t0)
        finally:
            os.unlink(path)
    return statistics.mean(times)


# ---------------------------------------------------------------------------
# DocxMint
# ---------------------------------------------------------------------------

def bench_docxmint(n_paragraphs: int) -> float:
    from docxmint import Document

    times = []
    for _ in range(_repeats(n_paragraphs)):
        path = _temp_path(".docx")
        try:
            t0 = time.perf_counter()
            with Document() as doc:
                for i in range(n_paragraphs):
                    doc.add_paragraph(
                        f"Paragraph {i}: The quick brown fox jumps over the lazy dog."
                    )
                n_rows = max(1, n_paragraphs // 10)
                table = doc.add_table(rows=n_rows, cols=3)
                for r in range(n_rows):
                    table[r, 0].text = f"Row {r}"
                    table[r, 1].text = "Value A"
                    table[r, 2].text = "Value B"
                doc.save(path)
            times.append(time.perf_counter() - t0)
        finally:
            os.unlink(path)
    return statistics.mean(times)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    results: dict[str, dict[str, float]] = {"python_docx": {}, "docxmint": {}}

    for size in DOCUMENT_SIZES:
        print(f"  python-docx  n={size:>6} ...", end=" ", flush=True)
        t = bench_python_docx(size)
        results["python_docx"][str(size)] = t
        print(f"{t:.3f}s")

        print(f"  docxmint     n={size:>6} ...", end=" ", flush=True)
        t = bench_docxmint(size)
        results["docxmint"][str(size)] = t
        print(f"{t:.3f}s")

    out = os.path.join(os.path.dirname(__file__), "benchmark_results.json")
    with open(out, "w") as f:
        json.dump(results, f, indent=2)
    print(f"\nResults written to {out}")
    print("Run scripts/generate_charts.py to produce the performance charts.")


if __name__ == "__main__":
    main()
