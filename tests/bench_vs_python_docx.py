"""Benchmark FastDOCX vs python-docx.

Run directly:
    python tests/bench_vs_python_docx.py

Measures wall-clock time for four workloads:
  - paragraphs   : 500 plain paragraphs
  - headings     : 200 headings across levels 1-3
  - table        : single 50x10 table with cell text
  - open+append  : open an existing 100-paragraph document and append 50 more

Each workload is repeated RUNS times; the median is reported.
"""

from __future__ import annotations

import os
import statistics
import sys
import tempfile
import time
from collections.abc import Callable

RUNS = 7
PARA_COUNT = 500
HEADING_COUNT = 200
TABLE_ROWS = 50
TABLE_COLS = 10
OPEN_EXISTING_PARAS = 100   # paragraphs in the pre-built fixture file
OPEN_APPEND_PARAS = 50      # paragraphs appended during the timed workload


# ---------------------------------------------------------------------------
# FastDOCX workloads
# ---------------------------------------------------------------------------


def _fastdocx_paragraphs(path: str) -> None:
    from fastdocx import Document

    doc = Document()
    for i in range(PARA_COUNT):
        doc.add_paragraph(f"Paragraph {i}", bold=i % 2 == 0)
    doc.save(path)


def _fastdocx_headings(path: str) -> None:
    from fastdocx import Document

    doc = Document()
    for i in range(HEADING_COUNT):
        doc.add_heading(f"Heading {i}", level=(i % 3) + 1)
    doc.save(path)


def _fastdocx_table(path: str) -> None:
    from fastdocx import Document

    doc = Document()
    table = doc.add_table(rows=TABLE_ROWS, cols=TABLE_COLS)
    for r in range(TABLE_ROWS):
        for c in range(TABLE_COLS):
            table[r, c].text = f"R{r}C{c}"
    doc.save(path)


def _fastdocx_open_append(source_path: str, out_path: str) -> None:
    from fastdocx import Document

    doc = Document(source_path)
    for i in range(OPEN_APPEND_PARAS):
        doc.add_paragraph(f"Appended {i}")
    doc.save(out_path)


# ---------------------------------------------------------------------------
# python-docx workloads
# ---------------------------------------------------------------------------


def _pydocx_paragraphs(path: str) -> None:
    import docx

    doc = docx.Document()
    for i in range(PARA_COUNT):
        p = doc.add_paragraph(f"Paragraph {i}")
        if i % 2 == 0:
            p.runs[0].bold = True
    doc.save(path)


def _pydocx_headings(path: str) -> None:
    import docx

    doc = docx.Document()
    for i in range(HEADING_COUNT):
        doc.add_heading(f"Heading {i}", level=(i % 3) + 1)
    doc.save(path)


def _pydocx_open_append(source_path: str, out_path: str) -> None:
    import docx

    doc = docx.Document(source_path)
    for i in range(OPEN_APPEND_PARAS):
        doc.add_paragraph(f"Appended {i}")
    doc.save(out_path)


def _pydocx_table(path: str) -> None:
    import docx

    doc = docx.Document()
    table = doc.add_table(rows=TABLE_ROWS, cols=TABLE_COLS)
    for r in range(TABLE_ROWS):
        for c in range(TABLE_COLS):
            table.cell(r, c).text = f"R{r}C{c}"
    doc.save(path)


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def _time_workload(fn: Callable[[str], None]) -> list[float]:
    times: list[float] = []
    for _ in range(RUNS):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as f:
            path = f.name
        t0 = time.perf_counter()
        fn(path)
        times.append(time.perf_counter() - t0)
    return times


def _time_open_workload(
    fn: Callable[[str, str], None], source_path: str
) -> list[float]:
    times: list[float] = []
    for _ in range(RUNS):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as f:
            out_path = f.name
        t0 = time.perf_counter()
        fn(source_path, out_path)
        times.append(time.perf_counter() - t0)
    return times


def _fmt(seconds: float) -> str:
    return f"{seconds * 1000:.1f} ms"


def _run_suite(
    label: str,
    fast_fn: Callable[[str], None],
    py_fn: Callable[[str], None],
) -> None:
    fast_times = _time_workload(fast_fn)
    py_times = _time_workload(py_fn)

    fast_med = statistics.median(fast_times)
    py_med = statistics.median(py_times)
    speedup = py_med / fast_med if fast_med > 0 else float("inf")

    print(f"  {label}")
    print(f"    fastdocx   : {_fmt(fast_med)}  (median of {RUNS})")
    print(f"    python-docx: {_fmt(py_med)}  (median of {RUNS})")
    print(f"    speedup    : {speedup:.1f}x")
    print()


def _run_open_suite(
    label: str,
    fast_fn: Callable[[str, str], None],
    py_fn: Callable[[str, str], None],
    source_path: str,
) -> None:
    fast_times = _time_open_workload(fast_fn, source_path)
    py_times = _time_open_workload(py_fn, source_path)

    fast_med = statistics.median(fast_times)
    py_med = statistics.median(py_times)
    speedup = py_med / fast_med if fast_med > 0 else float("inf")

    print(f"  {label}")
    print(f"    fastdocx   : {_fmt(fast_med)}  (median of {RUNS})")
    print(f"    python-docx: {_fmt(py_med)}  (median of {RUNS})")
    print(f"    speedup    : {speedup:.1f}x")
    print()


def _build_fixture() -> str:
    """Create a temporary .docx with OPEN_EXISTING_PARAS paragraphs for the open workload."""
    import docx

    doc = docx.Document()
    for i in range(OPEN_EXISTING_PARAS):
        doc.add_paragraph(f"Existing paragraph {i}")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc.save(path)
    return path


def main() -> None:
    # Check availability
    try:
        from fastdocx._native.loader import get_lib
        get_lib()
    except RuntimeError:
        print("ERROR: FastDOCX native binary not available — cannot benchmark.")
        sys.exit(1)

    try:
        import docx  # noqa: F401
    except ImportError:
        print("ERROR: python-docx not installed — run: pip install python-docx")
        sys.exit(1)

    print(f"Benchmark: FastDOCX vs python-docx  ({RUNS} runs each, median reported)\n")

    _run_suite(
        f"Paragraphs ({PARA_COUNT} paragraphs, alternating bold)",
        _fastdocx_paragraphs,
        _pydocx_paragraphs,
    )
    _run_suite(
        f"Headings ({HEADING_COUNT} headings, levels 1-3)",
        _fastdocx_headings,
        _pydocx_headings,
    )
    _run_suite(
        f"Table ({TABLE_ROWS}x{TABLE_COLS} with cell text)",
        _fastdocx_table,
        _pydocx_table,
    )

    fixture = _build_fixture()
    try:
        _run_open_suite(
            f"Open + append ({OPEN_EXISTING_PARAS} existing paragraphs, append {OPEN_APPEND_PARAS})",
            _fastdocx_open_append,
            _pydocx_open_append,
            fixture,
        )
    finally:
        os.unlink(fixture)


if __name__ == "__main__":
    main()
