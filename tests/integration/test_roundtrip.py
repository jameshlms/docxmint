"""Integration round-trip tests.

Generates a real .docx with DocxMint (requires the native binary), then opens
it with python-docx and asserts on paragraph text and table structure.

Automatically skipped when the native binary is not present.
"""

from __future__ import annotations

import os
import tempfile

import pytest

# ---------------------------------------------------------------------------
# Skip guard — import Document lazily inside each test so that importing this
# module never raises even when the binary is absent.
# ---------------------------------------------------------------------------

_SKIP_REASON = "Native DocxMint binary not available on this platform"


def _native_available() -> bool:
    """Return True only if the native handle loads without error."""
    try:
        from docxmint._native.handle import get_handle

        get_handle()
        return True
    except RuntimeError:
        return False


native_only = pytest.mark.skipif(not _native_available(), reason=_SKIP_REASON)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _open_docx(path: str) -> docx.Document:  # type: ignore[name-defined]  # noqa: F821
    import docx  # type: ignore[import-untyped]

    return docx.Document(path)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@native_only
def test_paragraph_roundtrip() -> None:
    """Paragraphs written by DocxMint are readable by python-docx."""
    from docxmint import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        doc.add_paragraph("Hello, world!")
        doc.add_paragraph().add_run("Bold text").bold = True
        doc.save(path)

        loaded = _open_docx(path)
        texts = [p.text for p in loaded.paragraphs]
        assert "Hello, world!" in texts
        assert "Bold text" in texts
    finally:
        os.unlink(path)


@native_only
def test_heading_roundtrip() -> None:
    """Headings written by DocxMint have the correct style in python-docx."""
    from docxmint import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        doc.add_heading("Chapter One", level=1)
        doc.add_heading("Section 1.1", level=2)
        doc.save(path)

        loaded = _open_docx(path)
        headings = {p.text: p.style.name for p in loaded.paragraphs if p.text}
        assert "Chapter One" in headings
        assert "Section 1.1" in headings
        assert "Heading 1" in headings["Chapter One"]
        assert "Heading 2" in headings["Section 1.1"]
    finally:
        os.unlink(path)


@native_only
def test_table_roundtrip() -> None:
    """Tables written by DocxMint have the correct cell text in python-docx."""
    from docxmint import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table[0, 0].text = "Region"
        table[0, 1].text = "Revenue"
        table[1, 0].text = "North"
        table[1, 1].text = "$1.2M"
        doc.save(path)

        loaded = _open_docx(path)
        assert len(loaded.tables) == 1
        tbl = loaded.tables[0]
        assert tbl.cell(0, 0).text == "Region"
        assert tbl.cell(0, 1).text == "Revenue"
        assert tbl.cell(1, 0).text == "North"
        assert tbl.cell(1, 1).text == "$1.2M"
    finally:
        os.unlink(path)


@native_only
def test_post_append_mutation_roundtrip() -> None:
    """Mutating a proxy after append (the _attach behavior) is reflected in the saved file."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        para = Paragraph("original")
        doc.paragraphs.append(para)
        para.text = "mutated after append"  # crosses FFI via LIVE proxy
        doc.save(path)

        loaded = _open_docx(path)
        texts = [p.text for p in loaded.paragraphs]
        assert "mutated after append" in texts
        assert "original" not in texts
    finally:
        os.unlink(path)


@native_only
def test_run_formatting_roundtrip() -> None:
    """Bold, italic, font size, and font name survive a save/load cycle."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph
    from docxmint.run import Run

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        para = Paragraph()
        doc.paragraphs.append(para)
        run = Run("styled text", bold=True, italic=True, font_size=14.0, font_name="Arial")
        para.runs.append(run)
        doc.save(path)

        loaded = _open_docx(path)
        runs = [r for p in loaded.paragraphs for r in p.runs if r.text == "styled text"]
        assert runs, "Expected run with 'styled text' not found"
        r = runs[0]
        assert r.bold is True
        assert r.italic is True
        assert r.font.name == "Arial"
        # python-docx stores font size in EMUs; 14pt = 177800
        assert r.font.size is not None
        assert abs(r.font.size.pt - 14.0) < 0.1
    finally:
        os.unlink(path)


@native_only
def test_post_append_run_mutation_roundtrip() -> None:
    """Mutating a Run proxy after append reflects in the saved file."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph
    from docxmint.run import Run

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        para = Paragraph()
        doc.paragraphs.append(para)
        run = Run("initial")
        para.runs.append(run)
        run.text = "updated"
        run.bold = True
        doc.save(path)

        loaded = _open_docx(path)
        runs = [r for p in loaded.paragraphs for r in p.runs if r.text == "updated"]
        assert runs, "Expected run with text 'updated' not found"
        assert runs[0].bold is True
    finally:
        os.unlink(path)


@native_only
def test_paragraph_alignment_roundtrip() -> None:
    """Paragraph alignment values survive a save/load cycle."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        doc.paragraphs.append(Paragraph("left"))
        doc.paragraphs.append(Paragraph("centered", alignment="center"))
        doc.paragraphs.append(Paragraph("right aligned", alignment="right"))
        doc.paragraphs.append(Paragraph("justified", alignment="justify"))
        doc.save(path)

        loaded = _open_docx(path)
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
        by_text = {p.text: p.alignment for p in loaded.paragraphs if p.text}
        assert by_text["centered"] == WD_ALIGN_PARAGRAPH.CENTER
        assert by_text["right aligned"] == WD_ALIGN_PARAGRAPH.RIGHT
        assert by_text["justified"] == WD_ALIGN_PARAGRAPH.JUSTIFY
    finally:
        os.unlink(path)


@native_only
def test_document_metadata_roundtrip() -> None:
    """Author and title document properties survive a save/load cycle."""
    from docxmint import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        doc.author = "Jane Smith"
        doc.title = "Q1 Report"
        doc.save(path)

        loaded = _open_docx(path)
        assert loaded.core_properties.author == "Jane Smith"
        assert loaded.core_properties.title == "Q1 Report"
    finally:
        os.unlink(path)


@native_only
def test_unicode_content_roundtrip() -> None:
    """Non-ASCII text survives a save/load cycle."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        texts = [
            "Héllo Wörld",          # Latin extended
            "日本語テスト",           # Japanese
            "Привет мир",           # Cyrillic
            "مرحبا بالعالم",        # Arabic
        ]
        for t in texts:
            doc.paragraphs.append(Paragraph(t))
        doc.save(path)

        loaded = _open_docx(path)
        loaded_texts = [p.text for p in loaded.paragraphs]
        for t in texts:
            assert t in loaded_texts, f"Missing: {t!r}"
    finally:
        os.unlink(path)


@native_only
def test_multiple_runs_in_paragraph_roundtrip() -> None:
    """Multiple runs in a single paragraph all survive a save/load cycle."""
    from docxmint import Document
    from docxmint.paragraph import Paragraph
    from docxmint.run import Run

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        para = Paragraph()
        doc.paragraphs.append(para)
        para.runs.append(Run("plain "))
        para.runs.append(Run("bold ", bold=True))
        para.runs.append(Run("italic", italic=True))
        doc.save(path)

        loaded = _open_docx(path)
        paras = [p for p in loaded.paragraphs if p.text.strip()]
        assert paras, "No non-empty paragraphs found"
        runs = paras[0].runs
        run_texts = [r.text for r in runs]
        assert "plain " in run_texts
        assert "bold " in run_texts
        assert "italic" in run_texts
        bold_run = next(r for r in runs if r.text == "bold ")
        italic_run = next(r for r in runs if r.text == "italic")
        assert bold_run.bold is True
        assert italic_run.italic is True
    finally:
        os.unlink(path)


@native_only
def test_open_nonexistent_raises_native_runtime_error() -> None:
    """Opening a file that does not exist raises NativeRuntimeError."""
    from docxmint import Document
    from docxmint.errors import NativeRuntimeError

    with pytest.raises(NativeRuntimeError):
        Document.open("/nonexistent_docxmint_test_xyz/missing.docx")


@native_only
def test_edit_nonexistent_raises_native_runtime_error() -> None:
    """Editing a file that does not exist raises NativeRuntimeError."""
    from docxmint import Document
    from docxmint.errors import NativeRuntimeError

    with pytest.raises(NativeRuntimeError):
        Document.edit("/nonexistent_docxmint_test_xyz/missing.docx")


@native_only
def test_save_to_bad_path_raises_native_runtime_error() -> None:
    """Saving to a path whose parent directory does not exist raises NativeRuntimeError."""
    from docxmint import Document
    from docxmint.errors import NativeRuntimeError

    doc = Document()
    try:
        with pytest.raises(NativeRuntimeError):
            doc.save("/nonexistent_docxmint_test_xyz/out.docx")
    finally:
        doc.close()


@native_only
def test_full_document_roundtrip() -> None:
    """The README quickstart example produces a valid document."""
    from docxmint import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = Document()
        doc.add_heading("Project Report", level=1)
        doc.add_paragraph().add_run("This report summarises Q1 findings.").bold = True

        table = doc.add_table(rows=3, cols=2)
        table[0, 0].text = "Region"
        table[0, 1].text = "Revenue"
        table[1, 0].text = "North"
        table[1, 1].text = "$1.2M"
        table[2, 0].text = "South"
        table[2, 1].text = "$0.9M"

        doc.save(path)

        loaded = _open_docx(path)
        para_texts = [p.text for p in loaded.paragraphs]
        assert "Project Report" in para_texts
        assert "This report summarises Q1 findings." in para_texts

        assert len(loaded.tables) == 1
        tbl = loaded.tables[0]
        assert tbl.cell(0, 0).text == "Region"
        assert tbl.cell(2, 1).text == "$0.9M"
    finally:
        os.unlink(path)
